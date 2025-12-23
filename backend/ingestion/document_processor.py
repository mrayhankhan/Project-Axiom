"""
Document processing pipeline for ML governance artifacts.
Handles text extraction, classification, and metadata extraction.
"""
from pathlib import Path
from typing import Dict, List, Optional
from datetime import datetime
import re
import PyPDF2
import docx
from pydantic import BaseModel


class DocumentMetadata(BaseModel):
    """Metadata for processed documents."""
    filename: str
    doc_type: str  # risk, bias, explainability, validation, assumptions
    model_name: Optional[str] = None
    version: Optional[str] = None
    date: Optional[str] = None
    file_size: int
    processed_at: str


class ProcessedDocument(BaseModel):
    """Processed document with extracted text and metadata."""
    content: str
    metadata: DocumentMetadata
    sections: List[Dict[str, str]]  # List of {title, content}


class DocumentProcessor:
    """Processes uploaded documents for ingestion into vector store."""
    
    # Classification keywords
    CLASSIFICATION_KEYWORDS = {
        "risk": ["risk", "threat", "vulnerability", "exposure", "mitigation"],
        "bias": ["bias", "fairness", "discrimination", "disparity", "equity", "protected"],
        "explainability": ["explainability", "interpretability", "shap", "lime", "feature importance"],
        "validation": ["validation", "testing", "performance", "accuracy", "metrics", "evaluation"],
        "assumptions": ["assumption", "limitation", "constraint", "dependency", "prerequisite"]
    }
    
    def __init__(self):
        pass
    
    def process_file(self, file_path: Path) -> ProcessedDocument:
        """Process a document file and extract content with metadata."""
        suffix = file_path.suffix.lower()
        
        if suffix == ".pdf":
            content = self._extract_pdf(file_path)
        elif suffix == ".docx":
            content = self._extract_docx(file_path)
        elif suffix in [".md", ".txt"]:
            content = self._extract_text(file_path)
        else:
            raise ValueError(f"Unsupported file type: {suffix}")
        
        # Extract sections
        sections = self._extract_sections(content)
        
        # Classify document type
        doc_type = self._classify_document(content)
        
        # Extract metadata from content
        model_name = self._extract_model_name(content)
        version = self._extract_version(content)
        date = self._extract_date(content)
        
        metadata = DocumentMetadata(
            filename=file_path.name,
            doc_type=doc_type,
            model_name=model_name,
            version=version,
            date=date,
            file_size=file_path.stat().st_size,
            processed_at=datetime.now().isoformat()
        )
        
        return ProcessedDocument(
            content=content,
            metadata=metadata,
            sections=sections
        )
    
    def _extract_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file."""
        text = []
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text.append(page.extract_text())
        return "\n".join(text)
    
    def _extract_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file."""
        doc = docx.Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])
    
    def _extract_text(self, file_path: Path) -> str:
        """Extract text from plain text or markdown file."""
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    
    def _extract_sections(self, content: str) -> List[Dict[str, str]]:
        """Extract sections from document based on headers."""
        sections = []
        
        # Pattern for markdown headers or numbered sections
        header_pattern = r'^#{1,3}\s+(.+?)$|^(\d+\.?\s+[A-Z].+?)$'
        
        lines = content.split('\n')
        current_section = None
        current_content = []
        
        for line in lines:
            match = re.match(header_pattern, line, re.MULTILINE)
            if match:
                # Save previous section
                if current_section:
                    sections.append({
                        "title": current_section,
                        "content": "\n".join(current_content).strip()
                    })
                
                # Start new section
                current_section = match.group(1) or match.group(2)
                current_content = []
            else:
                current_content.append(line)
        
        # Add last section
        if current_section:
            sections.append({
                "title": current_section,
                "content": "\n".join(current_content).strip()
            })
        
        # If no sections found, treat entire document as one section
        if not sections:
            sections.append({
                "title": "Document",
                "content": content
            })
        
        return sections
    
    def _classify_document(self, content: str) -> str:
        """Classify document type based on keyword frequency."""
        content_lower = content.lower()
        scores = {}
        
        for doc_type, keywords in self.CLASSIFICATION_KEYWORDS.items():
            score = sum(content_lower.count(keyword) for keyword in keywords)
            scores[doc_type] = score
        
        # Return type with highest score, default to 'risk'
        if max(scores.values()) > 0:
            return max(scores, key=scores.get)
        return "risk"
    
    def _extract_model_name(self, content: str) -> Optional[str]:
        """Extract model name from content."""
        # Look for patterns like "Model: XYZ" or "Model Name: XYZ"
        patterns = [
            r'Model\s*Name\s*[:：]\s*([^\n]+)',
            r'Model\s*[:：]\s*([^\n]+)',
            r'Algorithm\s*[:：]\s*([^\n]+)'
        ]
        
        for pattern in patterns:
            match = re.search(pattern, content, re.IGNORECASE)
            if match:
                return match.group(1).strip()
        
        return None
    
    def _extract_version(self, content: str) -> Optional[str]:
        """Extract version from content."""
        patterns = [
            r'Version\s*[:：]\s*([^\n]+)',
            r'v(\d+\.\d+\.?\d*)',
            r'Version\s+(\d+\.\d+\.?\d*)'
        ]
        
        for pattern in patterns:
            match = re.search(pattern, content, re.IGNORECASE)
            if match:
                return match.group(1).strip()
        
        return None
    
    def _extract_date(self, content: str) -> Optional[str]:
        """Extract date from content."""
        # Look for ISO dates or common date formats
        patterns = [
            r'Date\s*[:：]\s*(\d{4}-\d{2}-\d{2})',
            r'(\d{4}-\d{2}-\d{2})',
            r'Date\s*[:：]\s*([^\n]+)'
        ]
        
        for pattern in patterns:
            match = re.search(pattern, content, re.IGNORECASE)
            if match:
                return match.group(1).strip()
        
        return None
